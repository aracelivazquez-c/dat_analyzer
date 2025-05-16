"""
Módulo para procesar documentos Word y extraer datos estructurados
Soporta procesamiento de múltiples documentos y selección del documento más relevante
"""
import logging
import os
import re
from docx import Document
from difflib import SequenceMatcher

# Configurar logging
logger = logging.getLogger(__name__)

def get_document_content(file_path):
    """
    Extrae el contenido de texto de un documento Word
    
    Args:
        file_path (str): Ruta al documento Word
        
    Returns:
        str: Contenido de texto extraído del documento
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found at path: {file_path}")
        
        doc = Document(file_path)
        full_text = []
        
        # Extraer texto de párrafos
        for para in doc.paragraphs:
            if para.text.strip():  # Evitar añadir líneas vacías
                full_text.append(para.text)
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # Evitar añadir celdas vacías
                        row_text.append(cell.text.strip())
                if row_text:  # Evitar añadir filas vacías
                    full_text.append(" | ".join(row_text))
        
        logger.info(f"Successfully extracted content from document: {file_path}")
        return "\n".join(full_text)
    
    except Exception as e:
        logger.error(f"Error extracting content from Word document: {e}")
        raise

def get_document_data(file_path):
    """
    Procesa un documento Word y extrae datos estructurados
    para usar con KavakDataBot
    
    Args:
        file_path (str): Ruta al documento Word
        
    Returns:
        dict: Datos estructurados extraídos del documento
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found at path: {file_path}")
        
        # Obtener contenido completo
        content = get_document_content(file_path)
        
        # Procesar y estructurar el contenido
        doc = Document(file_path)
        
        # Extraer títulos y secciones
        sections = {}
        current_section = None
        section_content = []
        
        for para in doc.paragraphs:
            # Identificar encabezados basados en estilos o formato
            if para.style.name.startswith('Heading'):
                # Si tenemos una sección en proceso, guardarla
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content)
                    section_content = []
                
                # Iniciar nueva sección
                current_section = para.text.strip()
            elif current_section and para.text.strip():
                # Añadir contenido a la sección actual
                section_content.append(para.text.strip())
        
        # Guardar la última sección
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        # Extraer tablas
        tables_data = []
        for i, table in enumerate(doc.tables):
            table_data = {"table_id": i, "headers": [], "rows": []}
            
            # Extraer encabezados (primera fila)
            if table.rows:
                table_data["headers"] = [cell.text.strip() for cell in table.rows[0].cells]
                
                # Extraer datos (resto de filas)
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_data):  # Evitar filas vacías
                        table_data["rows"].append(row_data)
            
            tables_data.append(table_data)
        
        # Estructurar los datos para el bot
        structured_data = {
            "title": doc.paragraphs[0].text if doc.paragraphs else "Documentación Reporte | Compras MX",
            "raw_content": content,
            "sections": sections,
            "tables": tables_data
        }
        
        logger.info(f"Successfully processed document data from: {file_path}")
        return structured_data
    
    except Exception as e:
        logger.error(f"Error processing document data: {e}")
        raise
        
def process_multiple_documents(file_paths):
    """
    Procesa múltiples documentos Word y devuelve un diccionario con sus datos estructurados
    
    Args:
        file_paths (list): Lista de rutas a documentos Word
        
    Returns:
        dict: Diccionario donde las claves son las rutas de los documentos y los valores son los datos estructurados
    """
    documents_data = {}
    
    for file_path in file_paths:
        try:
            doc_data = get_document_data(file_path)
            # Usamos el nombre del archivo como identificador (sin la ruta completa ni la extensión)
            doc_id = os.path.splitext(os.path.basename(file_path))[0]
            # Aseguramos que el doc_id sea un string válido
            doc_id = str(doc_id).strip()
            documents_data[doc_id] = doc_data
            logger.info(f"Successfully processed document: {doc_id}")
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            # Continuamos con el siguiente documento en caso de error
            continue
    
    return documents_data

def calculate_similarity(text1, text2):
    """
    Calcula la similitud entre dos textos usando SequenceMatcher
    
    Args:
        text1 (str): Primer texto
        text2 (str): Segundo texto
        
    Returns:
        float: Puntuación de similitud entre 0 y 1
    """
    # Normalizamos los textos: minúsculas y eliminamos caracteres especiales
    text1 = re.sub(r'[^\w\s]', '', text1.lower())
    text2 = re.sub(r'[^\w\s]', '', text2.lower())
    
    # Para textos muy largos, comparamos solo el contenido relevante
    if len(text2) > 1000:
        # Encontrar palabras clave en la pregunta
        important_words = [word for word in text1.split() if len(word) > 3]
        
        # Buscar estas palabras en el documento y extraer contextos
        contexts = []
        for word in important_words:
            idx = text2.find(word)
            if idx != -1:
                # Extraer el contexto alrededor de la palabra
                start = max(0, idx - 100)
                end = min(len(text2), idx + 100)
                contexts.append(text2[start:end])
        
        # Si encontramos contextos, los combinamos para la comparación
        if contexts:
            text2 = " ".join(contexts)
    
    # Calculamos la similitud con el método de SequenceMatcher
    seq_similarity = SequenceMatcher(None, text1, text2).ratio()
    
    # También calculamos similitud basada en palabras clave
    words1 = set(text1.split())
    words2 = set(text2.split())
    
    # Palabras en común dividido por el total de palabras únicas en la pregunta
    if words1:
        keyword_similarity = len(words1.intersection(words2)) / len(words1)
    else:
        keyword_similarity = 0.0
    
    # Combinamos ambas métricas con más peso en las palabras clave
    combined_similarity = (seq_similarity * 0.4) + (keyword_similarity * 0.6)
    
    return combined_similarity

def find_most_relevant_document(question, documents_data):
    """
    Encuentra el documento más relevante para una pregunta dada
    
    Args:
        question (str): Pregunta del usuario
        documents_data (dict): Diccionario con los datos de los documentos
        
    Returns:
        tuple: (doc_id, doc_data) del documento más relevante
    """
    if not documents_data:
        raise ValueError("No hay documentos disponibles para consultar")
        
    # Si solo hay un documento, lo devolvemos directamente
    if len(documents_data) == 1:
        doc_id = list(documents_data.keys())[0]
        return doc_id, documents_data[doc_id]
    
    # Palabras clave importantes que pueden indicar documentos específicos
    keyword_mapping = {
        "originación": ["Tablero", "Originación", "Sales Mex"],
        "ventas": ["Performance", "Ventas", "Sales"],
        "sales": ["Performance", "Ventas", "Sales"],  # Añadimos "sales" como sinónimo de "ventas"
        "compras": ["Compras", "MX", "Supply"],
        "supply": ["Compras", "MX", "Supply"],  # Añadimos "supply" como sinónimo de "compras"
        "metrics": ["Metrics", "Supply", "Input"],
        "reservas": ["Atribución", "reservas"],
    }
    
    # Normalizamos la pregunta
    normalized_question = question.lower()
    
    # Comprobamos si la pregunta contiene palabras clave específicas
    priority_docs = []
    explicit_context = None  # Para guardar el contexto explícito (ventas o compras)
    
    # Lista de términos equivalentes para facilitar el reconocimiento
    equivalentes = {
        "sales": "ventas",
        "ventas": "ventas",
        "supply": "compras",
        "compras": "compras",
        "oportunidades de venta": "ventas",  # Frases contextuales comunes
        "funnel de ventas": "ventas",
        "funnel de compras": "compras",
        "oportunidad": "ventas",            # Términos que suelen estar asociados a ventas
        "oportunidades": "ventas",
        "venta": "ventas",
        "originacion": "ventas",
        "originación": "ventas",
        "leads": "ventas"
    }
    
    # Verificación especial para la frase "penetracion en oportunidades de venta"
    if "penetracion en oportunidades" in normalized_question or "penetración en oportunidades" in normalized_question:
        logger.info("FRASE ESPECÍFICA DETECTADA: penetracion en oportunidades de venta")
        explicit_context = "ventas"
        logger.info(f"CONTEXTO EXPLÍCITO FORZADO A VENTAS debido a la frase específica")
    
    # Procesar las equivalencias clave
    for palabra_clave, equivalente in equivalentes.items():
        if palabra_clave.lower() in normalized_question:
            logger.debug(f"Detectada palabra clave: {palabra_clave} (equivalente a {equivalente})")
            # Guardamos el contexto explícito si se menciona ventas o compras
            if equivalente in ["ventas", "compras"]:
                explicit_context = equivalente
                logger.info(f"CONTEXTO EXPLÍCITO DETECTADO: {explicit_context}")
            # Asegurarnos de que se detecte el término equivalente
            normalized_question = normalized_question.replace(palabra_clave.lower(), equivalente)
    
    # Ahora verificamos los topics en la pregunta normalizada
    for topic, keywords in keyword_mapping.items():
        # Verificar si el tema está en la pregunta
        if topic.lower() in normalized_question:
            logger.debug(f"Topic detectado en la pregunta: {topic}")
            # Agregamos documentos relacionados a la prioridad
            for doc_id in documents_data.keys():
                if any(keyword in doc_id for keyword in keywords):
                    # Si es un topic de ventas o compras, lo registramos para alta prioridad
                    if topic.lower() in ["ventas", "compras", "sales", "supply"]:
                        logger.debug(f"Topic importante: {topic} - alta prioridad para documento {doc_id}")
                    priority_docs.append(doc_id)
    
    # Documentos prioritarios únicos
    priority_docs = list(set(priority_docs))
    logger.debug(f"Documentos prioritarios basados en palabras clave: {priority_docs}")
    
    # Calculamos la similitud de la pregunta con cada documento
    similarities = {}
    for doc_id, doc_data in documents_data.items():
        # Combinamos título y contenido importante para la comparación
        title = doc_data.get("title", "")
        sections = doc_data.get("sections", {})
        
        # Concatenamos secciones importantes y primeras partes del contenido
        important_content = title + " "
        
        # Agregamos contenido de secciones importantes si existen
        for section_title, section_content in sections.items():
            important_content += section_title + " " + section_content[:200] + " "
        
        # Añadimos una muestra del contenido bruto 
        raw_content = doc_data.get("raw_content", "")
        sample_content = raw_content[:1000]  # Primeros 1000 caracteres
        
        doc_content = important_content + sample_content
        
        # Calculamos similitud 
        similarity = calculate_similarity(normalized_question, doc_content)
        
        # Bonus para documentos con palabras clave prioritarias
        if doc_id in priority_docs:
            similarity *= 1.5  # Aumentamos la relevancia en un 50%
        
        # ALTA PRIORIDAD para documentos que coincidan con el contexto explícito
        if explicit_context:
            # Si el contexto explícito es ventas, priorizamos documentos de ventas
            if explicit_context == "ventas":
                # Detección mejorada para el documento específico de Ventas
                if "ventas mx" in doc_id.lower():
                    logger.info(f"PRIORIDAD MÁXIMA para documento {doc_id} - Documento principal de ventas")
                    similarity *= 3.5  # Aumentamos la relevancia en un 250% adicional
                elif any(kw in doc_id.lower() for kw in ["ventas", "sales", "performance"]):
                    logger.info(f"ALTA PRIORIDAD para documento {doc_id} - Contexto de VENTAS explícito")
                    similarity *= 2.5  # Aumentamos la relevancia en un 150% adicional
            
            # Si el contexto explícito es compras, priorizamos documentos de compras
            elif explicit_context == "compras" and any(kw in doc_id.lower() for kw in ["compras", "supply", "metrics"]):
                logger.info(f"ALTA PRIORIDAD para documento {doc_id} - Contexto de COMPRAS explícito")
                similarity *= 2.5  # Aumentamos la relevancia en un 150% adicional
        
        # Prioridad específica para documentos de ventas con términos clave en la pregunta
        ventas_términos_clave = ["oportunidades", "handoff", "lead", "performance", "penetración"]
        if any(term in normalized_question for term in ventas_términos_clave) and "ventas mx" in doc_id.lower():
            logger.info(f"PRIORIDAD ESPECIAL para documento {doc_id} - Términos clave de ventas detectados")
            similarity *= 2.0  # Duplicamos la relevancia para términos clave específicos
            
        similarities[doc_id] = similarity
        logger.debug(f"Similitud con documento {doc_id}: {similarity:.4f}")
    
    # Ordenamos los documentos por similitud (de mayor a menor)
    sorted_docs = sorted(similarities.items(), key=lambda x: x[1], reverse=True)
    
    # Tomamos el documento con mayor similitud
    most_relevant_doc_id, max_similarity = sorted_docs[0]
    
    # Registramos información detallada para depuración
    logger.info(f"Documento más relevante para la pregunta: {most_relevant_doc_id} (similitud: {max_similarity:.4f})")
    logger.info(f"Top 3 documentos por similitud: {sorted_docs[:3]}")
    
    return most_relevant_doc_id, documents_data[most_relevant_doc_id]